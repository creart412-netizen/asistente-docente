import streamlit as st
import pdfplumber
import pytesseract
from pdf2image import convert_from_bytes
from PIL import Image
import language_tool_python
import docx
import pandas as pd
import openai
import os
from io import BytesIO
import tempfile
import datetime
import json
import uuid # Aunque no se usa directamente para la key del botón de navegación, es útil para otros IDs
import base64 # Importar para manejo de imágenes Base64
import random # Para mezclar las preguntas del juego
import streamlit.components.v1 as components # Importar para usar iframes

# Importar librerías para la conversión de documentos y generación de PDF
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx2pdf import convert as docx_to_pdf_convert

# ================= CONFIGURACIÓN DE LA APLICACIÓN =================
# CLAVE DE API DE OPENAI
# ¡IMPORTANTE! Para el DESPLIEGUE en Streamlit Community Cloud:
# 1. Asegúrate de tener un archivo .streamlit/secrets.toml con openai_api_key = "tu_clave_completa_aqui"
# 2. Las líneas de 'try-except' manejan esto.
#
# Para PRUEBAS LOCALES SIN secrets.toml:
# Puedes descomentar la línea 'openai.api_key = "sk-proj-..."' de abajo
# PERO ASEGÚRATE DE COMENTARLA ANTES DE SUBIR A GITHUB PARA DESPLIEGUE PÚBLICO.
try:
    openai.api_key = st.secrets["openai_api_key"]
except (AttributeError, KeyError):
    st.error("Error: La clave de API de OpenAI no se encontró en st.secrets.")
    st.warning("Para despliegue, crea un archivo `.streamlit/secrets.toml` con `openai_api_key = \"tu_clave_aqui\"`.")
    st.warning("Para pruebas **LOCALES**, puedes DESCOMENTAR la línea siguiente con tu clave directamente, PERO RECUERDA COMENTARLA ANTES DE SUBIR A GITHUB.")
    # --- DESCOMENTA LA LÍNEA DE ABAJO SOLO PARA PRUEBAS LOCALES SIN SECRETS.TOML ---
    # openai.api_key = "sk-proj-TU_CLAVE_VIEJA_O_DE_PRUEBA_AQUI_AQUI_AQUI"
    # --- FIN DEL BLOQUE DE PRUEBAS LOCALES ---
    if not hasattr(openai, 'api_key') or not openai.api_key:
        st.stop() # Detener la ejecución si la clave no está configurada


# LÍNEA DE DEPURACIÓN DE LA CLAVE API ELIMINADA.


# Configuración de Tesseract OCR
# ¡IMPORTANT! Para el DESPLIEGUE en Streamlit Community Cloud:
# COMENTA estas dos líneas para el despliegue. Streamlit Cloud instala tesseract-ocr
# como paquete de sistema (vía packages.txt) y pytesseract lo encontrará automáticamente.
#
# Para PRUEBAS LOCALES en tu PC:
# Mantén estas líneas DESCOMENTADAS y ASEGÚRATE de que las rutas sean las CORRECTAS
# donde instalaste Tesseract en tu computadora.
pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
os.environ["TESSDATA_PREFIX"] = r"C:\Program Files\Tesseract-OCR\tessdata"


TIPOS_ERROR = {
    "MORFOLOGIK_RULE_ES": "Ortográfico",
    "COMMA_PARENTHESIS_WHITESPACE": "Puntuación",
    "GRAMMAR": "Gramatical",
    "TYPO": "Ortográfico",
    "HUNSPELL_NO_SUGGESTION": "Ortográfico",
    "SPANISH_SPELLING_RULE": "Ortográfico",
    "MAYBE_UNNECESSARY_COMMA": "Puntuación",
}

# --- CSS Injection for Highlighting and specific button/expander styling ---
st.markdown("""
<style>
    /* Style for highlighted errors */
    .highlight-error {
        color: #DC3545; /* Red color for error */
        font-weight: bold;
        text-decoration: underline;
        background-color: #F8D7DA; /* Light red background for emphasis */
        padding: 2px 4px;
        border-radius: 3px;
    }
    /* Specific styling for the plus button to make it look nicer */
    div[data-testid="stVerticalBlock"] > div > div > div[data-testid="stColumn"] > button {
        background-color: #007BFF;
        color: white;
        border-radius: 50%; /* Make it round */
        width: 45px; /* Slightly larger */
        height: 45px; /* Slightly larger */
        font-size: 1.8em; /* Larger plus symbol */
        display: flex;
        align-items: center;
        justify-content: center;
        box-shadow: 3px 3px 8px rgba(0,0,0,0.25); /* More prominent shadow */
        transition: all 0.2s ease-in-out;
        border: none;
        margin-left: auto; /* Push to right */
        margin-right: 0;
    }
    div[data-testid="stVerticalBlock"] > div > div > div[data-testid="stColumn"] > button:hover {
        background-color: #0056B3;
        transform: scale(1.05);
        box-shadow: 4px 44px 10px rgba(0,0,0,0.3);
    }
    /* Styling for the file uploader within the chat (if still using expander for other purposes, like in chat history) */
    .stExpander {
        border: none !important;
        box-shadow: none !important;
    }
    .stExpander > div > div[data-testid="stVerticalBlock"] {
        border: 1px solid #D1D1D1;
        border-radius: 8px;
        padding: 10px;
        background-color: #FFFFFF;
        box-shadow: 1px 1px 3px rgba(0,0,0,0.05);
    }
    .stExpander > button[aria-expanded="false"] { /* Hide header of expander when collapsed */
        display: none;
    }
    .stExpander > button[aria-expanded="true"] { /* Style header of expander when open */
        background-color: #F0F2F6; /* Match background */
        color: #007BFF;
        font-weight: 600;
        border-radius: 8px;
        padding: 5px 10px;
        border: 1px solid #E0E0E0;
    }
    /* Ensure chat input is visible */
    [data-testid="stTextInput"] div.st-af { /* Targets the input div directly below stTextInput */
        background-color: #FFFFFF; /* Ensure input field is white */
        border-radius: 8px;
        border: 1px solid #D1D1D1;
        box-shadow: 1px 1px 3px rgba(0,0,0,0.05);
        padding: 5px; /* Adjust padding if needed */
    }
     /* Styling for the file management list */
    /* Adjusting specific download/delete buttons in "Mis Archivos Guardados" */
    .stDownloadButton button { /* Default Streamlit download button */
        background-color: #28A745; /* Green for download */
        color: white;
        border-radius: 5px;
        border: none;
        padding: 8px 12px;
        font-size: 0.9em;
        margin: 2px;
        transition: background-color 0.2s ease-in-out, transform 0.1s ease-in-out;
        box-shadow: 1px 1px 3px rgba(0,0,0,0.1);
    }
    .stDownloadButton button:hover {
        background-color: #218838;
        transform: translateY(-1px);
    }
    /* Targeting the delete button specifically by Streamlit's internal test ID if possible, or context */
    .stButton button[kind="secondary"] { /* Re-override to ensure delete button is red if it falls into this class */
        background-color: #DC3545 !important; /* Red for delete */
        color: white !important;
        border-radius: 5px !important;
        border: none !important;
        padding: 8px 12px !important;
        font-size: 0.9em !important;
        margin: 2px !important;
        transition: background-color 0.2s ease-in-out, transform 0.1s ease-in-out !important;
        box_shadow: 1px 1px 3px rgba(0,0,0,0.1) !important;
    }
    .stButton button[kind="secondary"]:hover {
        background-color: #C82333 !important;
        transform: translateY(-1px) !important;
    }
    /* Specific styling for the columns in Mis Archivos Guardados */
    div[data-testid^="stHorizontalBlock"] > div[data-testid^="stVerticalBlock"] {
        padding: 5px;
    }
    div[data-testid^="stHorizontalBlock"] {
        border-bottom: 1px solid #EEEEEE; /* Light separator for list items */
        padding-bottom: 10px;
        margin-bottom: 10px;
    }
    /* Custom styles for navigation buttons (these apply to the horizontal buttons if used) */
    div[data-testid="stHorizontalBlock"] button {
        background-color: #007BFF;
        color: white;
        border-radius: 8px;
        padding: 10px 15px;
        font-size: 1em;
        margin: 5px;
        border: none;
        box-shadow: 2px 2px 5px rgba(0,0,0,0.2);
        transition: all 0.2s ease-in-out;
    }
    div[data-testid="stHorizontalBlock"] button:hover {
        background-color: #0056B3;
        transform: translateY(-2px);
    }
    div[data-testid="stHorizontalBlock"] button[aria-pressed="true"] { /* Style for active/selected button */
        background-color: #0056B3;
        border: 2px solid #00BFFF; /* Slightly lighter blue border for active state */
        transform: translateY(0);
        box-shadow: inset 1px 1px 3px rgba(0,0,0,0.3);
    }
</style>
""", unsafe_allow_html=True)

# Configuración de la página: Importante para el layout
st.set_page_config(page_title="Asistente Docente", page_icon="📚", layout="wide")

# ==================== CONFIGURACIÓN DE ALMACENAMIENTO LOCAL ====================
LOCAL_FILES_DIR = "generated_files"
if not os.path.exists(LOCAL_FILES_DIR):
    os.makedirs(LOCAL_FILES_DIR)

METADATA_FILE = os.path.join(LOCAL_FILES_DIR, "files_metadata.json")

@st.cache_resource
def load_metadata():
    if os.path.exists(METADATA_FILE):
        try:
            with open(METADATA_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        except json.JSONDecodeError:
            st.error("Error al leer el archivo de metadatos. Se iniciará un historial vacío.")
            return []
        except Exception as e:
            st.error(f"Error al cargar metadatos: {e}. Se iniciará un historial vacío.")
            return []
    return []

def save_metadata(metadata_list):
    try:
        with open(METADATA_FILE, 'w', encoding='utf-8') as f:
            json.dump(metadata_list, f, indent=4)
    except Exception as e:
        st.error(f"Error al guardar los metadatos: {e}")

if 'local_files_history' not in st.session_state:
    st.session_state.local_files_history = load_metadata()

# Función para guardar un archivo directamente en el disco local y actualizar metadatos
def save_file_to_local_disk_and_update_metadata(file_bytes, original_file_base_name, output_ext, file_type):
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = str(uuid.uuid4())
    unique_file_name = f"{original_file_base_name}_{timestamp}_{unique_id}{output_ext}"
    file_path = os.path.join(LOCAL_FILES_DIR, unique_file_name)

    try:
        with open(file_path, 'wb') as f:
            f.write(file_bytes)

        metadata_entry = {
            "id": unique_id,
            "name": unique_file_name,
            "type": file_type,
            "generated_at": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "local_path": file_path
        }
        st.session_state.local_files_history.append(metadata_entry)
        save_metadata(st.session_state.local_files_history)
        st.success(f"'{unique_file_name}' guardado automáticamente en la carpeta local.")
        return True
    except Exception as e:
        st.error(f"Error al guardar el archivo localmente: {e}")
        return False

def delete_local_file(file_id):
    updated_history = []
    deleted_file_name = ""
    file_path_to_delete = ""
    for entry in st.session_state.local_files_history:
        if entry["id"] == file_id:
            file_path_to_delete = entry["local_path"]
            deleted_file_name = entry["name"]
        else:
            updated_history.append(entry)

    if file_path_to_delete and os.path.exists(file_path_to_delete):
        try:
            os.remove(file_path_to_delete)
            st.success(f"Archivo físico '{deleted_file_name}' eliminado.")
        except OSError as e:
            st.error(f"Error al eliminar el archivo físico del disco '{deleted_file_name}': {e}. Es posible que necesites permisos.")
            updated_history.append([entry for entry in st.session_state.local_files_history if entry["id"] == file_id][0])

    st.session_state.local_files_history = updated_history
    save_metadata(st.session_state.local_files_history)
    st.rerun()

# Función para convertir un archivo subido a base64
def get_file_base64_and_mime_type(uploaded_file):
    if uploaded_file is not None:
        bytes_data = uploaded_file.getvalue()
        base64_encoded_data = base64.b64encode(bytes_data).decode('utf-8')
        return base64_encoded_data, uploaded_file.type
    return None, None

# Función para extraer texto de documentos (PDF y DOCX)
def extract_text_from_document(uploaded_file):
    text_content = ""
    # Ensure the file pointer is at the beginning
    uploaded_file.seek(0)
    if uploaded_file.name.lower().endswith(".pdf"):
        try:
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    text_content += page.extract_text() or ""
            if not text_content.strip(): # Fallback to OCR if no digital text
                # Reset file pointer for convert_from_bytes
                uploaded_file.seek(0)
                pdf_bytes = uploaded_file.getvalue()
                images = convert_from_bytes(pdf_bytes)
                for i, image in enumerate(images):
                    st.text(f"Procesando página {i+1} con OCR...")
                    text_content += pytesseract.image_to_string(image, lang="spa")
        except Exception as e:
            st.warning(f"Error al extraer texto del PDF: {e}. Se usará texto vacío.")
            text_content = ""
    elif uploaded_file.name.lower().endswith(".docx"):
        try:
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                text_content += para.text + "\n"
        except Exception as e:
            st.warning(f"Error al extraer texto del DOCX: {e}. Se usará texto vacío.")
            text_content = ""
    return text_content

# Funciones de parseo y formateo para el contenido generado por IA
def parse_and_format_exam(raw_text):
    formatted_output = []
    lines = raw_text.split('\n')
    current_block = {"type": "paragraph", "content": []}

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            # New paragraph/block separator
            if current_block["content"]:
                formatted_output.append(current_block)
            current_block = {"type": "paragraph", "content": []}
            continue

        if line_stripped.lower().startswith("pregunta"):
            if current_block["content"]:
                formatted_output.append(current_block)
            current_block = {"type": "question_answer", "question": line_stripped, "answer": ""}
        elif line_stripped.lower().startswith("respuesta") and current_block["type"] == "question_answer":
            current_block["answer"] = line_stripped
        elif line_stripped.lower().startswith("respuesta") and current_block["type"] == "question_answer": # Typo, this line is redundant
            current_block["answer"] = line_stripped
        elif current_block["type"] == "question_answer":
            if current_block["answer"]: # If answer has started, append to answer
                current_block["answer"] += "\n" + line_stripped
            else: # Otherwise, still part of question
                current_block["question"] += "\n" + line_stripped
        else:
            current_block["content"].append(line_stripped)

    if current_block["content"] or (current_block["type"] == "question_answer" and current_block["question"]):
        formatted_output.append(current_block)

    return formatted_output

def parse_and_format_lesson_plan(raw_text):
    formatted_output = []
    lines = raw_text.split('\n')
    current_block = None

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            if current_block:
                formatted_output.append(current_block)
                current_block = None
            continue

        # Check for Week headers (more robustly, e.g., "Semana X")
        if line_stripped.lower().startswith("semana"):
            if current_block:
                formatted_output.append(current_block)
            current_block = {"type": "week", "title": line_stripped, "sections": {}}
        elif current_block and current_block["type"] == "week":
            # Check for section headers within a week
            section_found = False
            for keyword in ["objetivos:", "actividades de enseñanza-aprendizaje:", "recursos sugeridos:", "métodos de evaluación:"]:
                if line_stripped.lower().startswith(keyword):
                    section_name = keyword.replace(":", "").title() # "Objetivos", "Actividades De Enseñanza-Aprendizaje
                    if section_name not in current_block["sections"]:
                        current_block["sections"][section_name] = []
                    current_block["sections"][section_name].append(line_stripped[len(keyword):].strip()) # Add remaining text to section
                    section_found = True
                    break

            if not section_found:
                # If no specific section keyword, append to a general 'content' within the week or last section
                if "content" not in current_block:
                    current_block["content"] = []
                current_block["content"].append(line_stripped)
        else: # General paragraph outside any specific structure
            if not current_block or current_block["type"] != "paragraph":
                current_block = {"type": "paragraph", "content": []}
            current_block["content"].append(line_stripped) # Corrected line: 'current_content' changed to 'current_block["content"]'

    if current_block:
        formatted_output.append(current_block)

    return formatted_output

def parse_and_format_math_problem(raw_text):
    """
    Parses the raw text from AI into problem, steps, and final answer.
    Assumes AI output is structured with clear headers like:
    "Problema:", "Pasos:", "Respuesta Final:".
    """
    problem = ""
    steps = []
    answer = ""

    lines = raw_text.split('\n')
    current_section = None

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        if line_stripped.startswith("Problema:"):
            current_section = "problem"
            problem = line_stripped[len("Problema:"):].strip()
        elif line_stripped.startswith("Pasos:"):
            current_section = "steps"
        elif line_stripped.startswith("Respuesta Final:"):
            current_section = "answer"
            answer = line_stripped[len("Respuesta Final:"):].strip()
        elif current_section == "problem":
            problem += "\n" + line_stripped
        elif current_section == "steps":
            steps.append(line_stripped)
        elif current_section == "answer":
            answer += "\n" + line_stripped # In case answer spans multiple lines

    return {"problem": problem.strip(), "steps": steps, "answer": answer.strip()}

def generate_pdf_from_math_problem(parsed_math_problem_data):
    """
    Generates a PDF from the parsed math problem data.
    """
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()

    story = []

    # Title Style
    title_style = ParagraphStyle(
        'Title',
        parent=styles['h1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        alignment=TA_CENTER,
        spaceAfter=12
    )
    story.append(Paragraph("Problema de Matemáticas con Solución Paso a Paso", title_style))
    story.append(Spacer(1, 0.2 * inch))

    # Problem Section
    story.append(Paragraph("<b>Problema:</b>", styles['h2']))
    # Render problem with line breaks
    problem_text = parsed_math_problem_data['problem'].replace('\n', '<br/>')
    story.append(Paragraph(problem_text, styles['Normal']))
    story.append(Spacer(1, 0.2 * inch))

    # Steps Section
    story.append(Paragraph("<b>Solución Paso a Paso:</b>", styles['h2']))
    for step in parsed_math_problem_data['steps']:
        if step.strip(): # Only add non-empty steps
            # Render step with line breaks
            step_text = step.replace('\n', '<br/>')
            story.append(Paragraph(step_text, styles['Normal']))
            story.append(Spacer(1, 0.1 * inch))
    story.append(Spacer(1, 0.2 * inch))

    # Final Answer Section
    story.append(Paragraph("<b>Respuesta Final:</b>", styles['h2']))
    # Render answer with line breaks
    answer_text = parsed_math_problem_data['answer'].replace('\n', '<br/>')
    story.append(Paragraph(answer_text, styles['Normal']))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()


def highlight_errors_in_text(text, matches):
    """
    Highlights errors in the given text using HTML spans.

    Args:
        text (str): The original text.
        matches (list): A list of Match objects from language_tool_python.

    Returns:
        str: HTML string with errors highlighted.
    """
    if not matches:
        return text.replace('\n', '<br>') # Convert newlines for HTML display

    # Sort matches by offset to apply highlighting without disrupting subsequent offsets
    # Process from the end of the text to avoid index issues with string manipulation
    sorted_matches = sorted(matches, key=lambda x: x.offset, reverse=True)

    highlighted_text = list(text) # Convert string to list of characters for easier manipulation

    for match in sorted_matches:
        start = match.offset
        end = match.offset + match.errorLength

        if start < len(highlighted_text) and end <= len(highlighted_text):
            # Extract the error phrase
            error_phrase = "".join(highlighted_text[start:end])

            # Create the highlighted span
            highlighted_span = f'<span class="highlight-error">{error_phrase}</span>'

            # Replace the original characters with the HTML span
            highlighted_text[start:end] = list(highlighted_span)

    return "".join(highlighted_text).replace('\n', '<br>')


# ==================== NAVEGACIÓN (Diseño Horizontal Minimalista) ====================

# Definimos las secciones y sus iconos
tabs_names = [
    "✏️ Revisión Ortográfica",
    "📂 Contenido Educativo",
    "🔢 Problemas Matemáticos",
    "💬 Chat Educativo",
    "🔄 Convertidor de Docs",
    "🎮 Juegos Educativos",
    "🗄️ Mis Archivos"
]

# Logo y título principal en la parte superior del cuerpo de la aplicación
try:
    st.image("assets/logo.png", width=150)
except FileNotFoundError:
    st.warning("No se encontró 'assets/logo.png'. Asegúrate de que la imagen exista en la ruta correcta.")
st.title("📚 Asistente Docente")
st.markdown("---") # Separador visual

# Inicializamos la sección seleccionada en el estado de sesión
if "selected_section" not in st.session_state:
    st.session_state.selected_section = tabs_names[0] # Selecciona la primera por defecto

# Creamos una fila de columnas para los botones de navegación
# Usamos un número fijo de columnas para asegurar el espaciado
# Puedes ajustar 'num_columns' para cambiar cuántos botones caben en una fila antes de que se envuelvan.
num_columns = 4
cols = st.columns(num_columns)

# Generar los botones de navegación dentro de las columnas
for i, tab_name in enumerate(tabs_names):
    with cols[i % num_columns]: # Usamos el módulo para ciclar entre las columnas
        # Aquí está la CORRECCIÓN: la 'key' debe ser estable y única por el nombre del botón.
        # Ya NO usamos uuid.uuid4() aquí, ya que causaba que el botón perdiera su estado de click.
        button_key = f"nav_button_{tab_name.replace(' ', '_').lower()}"

        # Si el nombre del botón es el mismo que la sección seleccionada, agregamos un estilo visual
        is_selected = (st.session_state.selected_section == tab_name)
        button_style_class = "selected-button" if is_selected else ""

        if st.button(tab_name, key=button_key, help=f"Ir a la sección {tab_name}"):
            # Solo actualiza la sesión y haz rerun si la sección realmente cambió
            if st.session_state.selected_section != tab_name:
                st.session_state.selected_section = tab_name
                st.rerun() # Forzar un re-run para cambiar la página principal


st.markdown("---") # Separador visual después de los botones de navegación


# ==================== CONTENIDO DE LAS SECCIONES ====================

# Usamos if/elif para mostrar el contenido de la sección seleccionada
if st.session_state.selected_section == "✏️ Revisión Ortográfica":
    st.header("Revisión de Textos")
    archivo = st.file_uploader("Sube un archivo PDF o Word", type=["pdf", "docx"])

    if 'last_extracted_text' not in st.session_state:
        st.session_state.last_extracted_text = ""
    if 'last_found_matches' not in st.session_state:
        st.session_state.last_found_matches = []

    # Inicializar la herramienta de lenguaje aquí para que se cachee
    @st.cache_resource
    def get_language_tool():
        try:
            # FIX FINAL: Forzar el uso del servidor remoto para evitar la dependencia de Java local
            return language_tool_python.LanguageTool(language='es-ES', remote_server_url='https://languagetool.org/api/v2/')
        except Exception as e:
            st.error(f"Error al inicializar LanguageTool: {e}. Asegúrate de tener conexión a internet.")
            st.stop()
    
    tool = get_language_tool()

    if archivo is not None:
        texto_extraido = ""
        archivo_name_lower = archivo.name.lower()
        # Ensure the file pointer is at the beginning
        archivo.seek(0)

        if archivo_name_lower.endswith(".pdf"):
            try:
                with pdfplumber.open(archivo) as pdf:
                    for pagina in pdf.pages:
                        texto_extraido += pagina.extract_text() or ""
            except Exception as e:
                st.error(f"Error al intentar leer texto digital del PDF: {e}")
                pass

            if not texto_extraido.strip():
                st.warning("No se encontró texto digital, usando OCR (esto puede tardar)...")
                try:
                    archivo.seek(0) # Reset pointer before reading bytes for OCR
                    pdf_bytes = archivo.getvalue()
                    images = convert_from_bytes(pdf_bytes)
                    for i, image in enumerate(images):
                        st.text(f"Procesando página {i+1} con OCR...")
                        texto_extraido += pytesseract.image_to_string(image, lang="spa")
                    st.success("OCR completado.")
                except Exception as e:
                    st.error(f"Error al procesar el PDF con OCR: {e}. Asegúrate de que Tesseract esté bien instalado y configurado.")
                    texto_extraido = ""

        elif archivo_name_lower.endswith(".docx"):
            try:
                doc = docx.Document(archivo)
                for parrafo in doc.paragraphs:
                    texto_extraido += parrafo.text + "\n"
            except Exception as e:
                st.error(f"Error al leer el archivo Word: {e}")
                texto_extraido = ""

        # Update session state with the new extracted text
        st.session_state.last_extracted_text = texto_extraido
        st.session_state.last_found_matches = [] # Clear previous matches

    # Check if we have text to process (either newly uploaded or from previous session)
    if st.session_state.last_extracted_text.strip():
        # Display the extracted text (or the one from session state)
        st.subheader("Texto extraído:")
        st.markdown(f"<div style='border: 1px solid #ddd; padding: 10px; border-radius: 5px; overflow-wrap: break-word;'>{st.session_state.last_extracted_text.replace('\n', '<br>')}</div>", unsafe_allow_html=True)

        if st.button("Revisar Texto", key="revisar_texto_btn"):
            with st.spinner("⏳ Realizando revisión ortográfica y gramatical..."):
                try:
                    # 'tool' ya está inicializado arriba y cacheado
                    matches = tool.check(st.session_state.last_extracted_text)
                    st.session_state.last_found_matches = matches # Store matches in session state

                    if matches:
                        st.subheader("Texto Corregido y Resaltado:")
                        highlighted_html = highlight_errors_in_text(st.session_state.last_extracted_text, matches)
                        st.markdown(f"<div style='border: 1px solid #ddd; padding: 10px; border-radius: 5px; overflow-wrap: break-word;'>{highlighted_html}</div>", unsafe_allow_html=True)
                        st.success("🎉 ¡Revisión completada! Errores resaltados en rojo.")

                    else:
                        st.success("🎉 ¡No se encontraron errores ortográficos o gramaticales! Tu texto está impecable.")
                except Exception as e:
                    st.error(f"Ocurrió un error al realizar la revisión: {e}. Por favor, inténtalo de nuevo más tarde.")

        # If there are already matches in session state (from a previous run), display them
        elif st.session_state.last_found_matches:
            st.subheader("Resultado de la Última Revisión:")
            highlighted_html = highlight_errors_in_text(st.session_state.last_extracted_text, st.session_state.last_found_matches)
            st.markdown(f"<div style='border: 1px solid #ddd; padding: 10px; border-radius: 5px; overflow-wrap: break-word;'>{highlighted_html}</div>", unsafe_allow_html=True)
            st.info("Revisión previamente realizada. Haz clic en 'Revisar Texto' para volver a procesar.")

    else:
        st.info("Sube un archivo (PDF o Word) para iniciar la revisión ortográfica y gramatical.")


elif st.session_state.selected_section == "📂 Contenido Educativo":
    st.header("Generador de Contenido Educativo (Exámenes / Planes de Clases)")
    archivo_tema = st.file_uploader("Sube el temario (PDF o Word) para generar contenido", type=["pdf", "docx"], key="temario")

    # Almacenar el contenido generado en el estado de sesión para que no desaparezca
    if 'generated_doc_display_content' not in st.session_state:
        st.session_state.generated_doc_display_content = None
    if 'generated_doc_display_name' not in st.session_state:
        st.session_state.generated_doc_display_name = ""
    if 'generated_doc_display_mime' not in st.session_state:
        st.session_state.generated_doc_display_mime = ""
    if 'generated_doc_display_type' not in st.session_state:
        st.session_state.generated_doc_display_type = ""
    if 'raw_generated_text_for_display' not in st.session_state:
        st.session_state.raw_generated_text_for_display = ""

    if archivo_tema is not None:
        texto_tema = ""
        # Seek to the beginning of the file before processing
        archivo_tema.seek(0)
        if archivo_tema.name.lower().endswith(".pdf"):
            try:
                with pdfplumber.open(archivo_tema) as pdf:
                    for pagina in pdf.pages:
                        texto_tema += pagina.extract_text() or ""
                if not texto_tema.strip(): # Fallback to OCR if no digital text
                    archivo_tema.seek(0) # Reset pointer
                    pdf_bytes = archivo_tema.getvalue()
                    images = convert_from_bytes(pdf_bytes)
                    for i, image in enumerate(images):
                        texto_tema += pytesseract.image_to_string(image, lang="spa")
            except Exception as e:
                st.error(f"Error al leer el PDF del temario: {e}. Intenta con otro archivo o verifica el formato.")
        elif archivo_tema.name.lower().endswith(".docx"):
            try:
                doc = docx.Document(archivo_tema)
                for parrafo in doc.paragraphs:
                    texto_tema += parrafo.text + "\n"
            except Exception as e:
                st.error(f"Error al leer el archivo Word del temario: {e}.")

        if texto_tema.strip():
            st.success("✅ Temario cargado con éxito.")
            nivel = st.selectbox("Selecciona el nivel educativo", ["Preescolar", "Primaria", "Secundaria", "Universidad"], key="nivel_generador")
            # --- MODIFICACIÓN AQUÍ: Añadir "(mensual)" al "Plan de Clases" ---
            opcion = st.radio("¿Qué quieres generar?", ["Examen", "Plan de Clases (mensual)"], key="tipo_generacion")
            # --- FIN DE LA MODIFICACIÓN ---
            instrucciones_extra = st.text_area("Notas para la IA (opcional)", placeholder="Ejemplo: hazlo resumido, incluye ejemplos, etc.", key="instrucciones_generador")

            if st.button("Generar contenido", key="boton_generar"):
                if not openai.api_key:
                    st.error("La clave de API de OpenAI no está configurada. Por favor, revisa la configuración del script.")
                else:
                    with st.spinner("⏳ Generando contenido con IA... Esto puede tardar unos segundos."):
                        contenido_texto_ia = ""
                        try:
                            if opcion == "Examen":
                                prompt = f"Eres un docente experto en nivel {nivel}. A partir del siguiente temario, genera 10 preguntas variadas (opción múltiple, verdadero/falso, y abiertas) con sus respuestas correctas. Asegúrate de que el examen sea coherente y esté bien estructurado para el nivel. Utiliza un lenguaje apropiado para el nivel educativo.\n\nTemario:\n{texto_tema}\n\nInstrucciones adicionales: {instrucciones_extra if instrucciones_extra else 'Ninguna'}"
                            # --- MODIFICACIÓN AQUÍ: Adaptar la lógica para el nuevo nombre "Plan de Clases (mensual)" ---
                            elif opcion == "Plan de Clases (mensual)": # Antes era "Plan de Clases"
                                prompt = f"Eres un docente experto en nivel {nivel}. A partir del siguiente temario, genera un plan de clases semanal detallado para un curso de 4 semanas. Cada semana debe incluir objetivos claros, actividades de enseñanza-aprendizaje, recursos sugeridos y métodos de evaluación. Adapta el lenguaje y las actividades al nivel educativo especificado. Considera un formato claro y fácil de leer.\n\nTemario:\n{texto_tema}\n\nInstrucciones adicionales: {instrucciones_extra if instrucciones_extra else 'Ninguna'}"
                            # --- FIN DE LA MODIFICACIÓN ---

                            respuesta = openai.chat.completions.create(
                                model="gpt-4o-mini",
                                messages=[{"role": "user", "content": prompt}],
                                temperature=0.7
                            )
                            contenido_texto_ia = respuesta.choices[0].message.content

                            # Convertir el texto de la IA a DOCX bytes
                            doc = docx.Document()
                            for para_text in contenido_texto_ia.split('\n'):
                                if para_text.strip():
                                    doc.add_paragraph(para_text)

                            doc_buffer = BytesIO()
                            doc.save(doc_buffer)
                            doc_buffer.seek(0)
                            docx_bytes = doc_buffer.getvalue()

                            # Guardar inmediatamente en el disco local y actualizar metadatos
                            # Asegúrate de que el `file_type` refleje el nuevo nombre si es necesario para metadatos
                            base_name_for_file = f"{opcion.replace(' ', '_').replace('(', '').replace(')', '').lower()}_{nivel.replace(' ', '_')}"
                            if save_file_to_local_disk_and_update_metadata(docx_bytes, base_name_for_file, ".docx", opcion.lower()):
                                # Si se guardó localmente con éxito, actualizamos session_state para la descarga y display
                                st.session_state.generated_doc_display_content = docx_bytes
                                st.session_state.generated_doc_display_name = f"{base_name_for_file}.docx"
                                st.session_state.generated_doc_display_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                st.session_state.generated_doc_display_type = opcion.lower()
                                st.session_state.raw_generated_text_for_display = contenido_texto_ia # Guardar texto original para vista previa
                                st.info("Puedes descargar el archivo o revisarlo en 'Mis Archivos Guardados (Local)'.")
                            else:
                                st.error("No se pudo guardar el archivo localmente. La descarga no estará disponible.")
                                st.session_state.generated_doc_display_content = None
                                st.session_state.raw_generated_text_for_display = ""

                        except openai.RateLimitError:
                            st.error("🚫 ¡Error de API! Se han agotado los tokens de OpenAI. Por favor, verifica tu plan y facturación en platform.openai.com/usage.")
                            st.session_state.generated_doc_display_content = None
                            st.session_state.raw_generated_text_for_display = ""
                        except openai.APIError as e:
                            st.error(f"🚫 Error de la API de OpenAI: {e}. Por favor, inténtalo de nuevo más tarde.")
                            st.session_state.generated_doc_display_content = None
                            st.session_state.raw_generated_text_for_display = ""
                        except Exception as e:
                            st.error(f"Ocurrió un error inesperado al generar el contenido: {e}. Por favor, revisa el temario e inténtalo de nuevo.")
                            st.session_state.generated_doc_display_content = None
                            st.session_state.raw_generated_text_for_display = ""

                        st.rerun()

            # Mostrar el contenido y botón de descarga si existe contenido generado/convertido
            if st.session_state.generated_doc_display_content is not None:
                st.markdown("### 📄 Resultado Generado")

                # Formatear y mostrar el contenido basado en el tipo
                if st.session_state.generated_doc_display_type == "examen":
                    parsed_content = parse_and_format_exam(st.session_state.raw_generated_text_for_display)
                    for block in parsed_content:
                        if block["type"] == "question_answer":
                            st.markdown(f"**{block['question']}**")
                            with st.expander("Ver Respuesta"):
                                st.write(block["answer"])
                        elif block["type"] == "paragraph" and block["content"]:
                            st.write("\n".join(block["content"]))
                # --- MODIFICACIÓN AQUÍ: Adaptar la lógica para el nuevo nombre "Plan de Clases (mensual)" ---
                elif st.session_state.generated_doc_display_type == "plan de clases (mensual)": # Antes era "plan de clases"
                    parsed_content = parse_and_format_lesson_plan(st.session_state.raw_generated_text_for_display)
                    for block in parsed_content:
                        if block["type"] == "week":
                            st.markdown(f"## {block['title']}")
                            if "content" in block and block["content"]:
                                st.write("\n".join(block["content"]))
                            for section_name, section_content_list in block["sections"].items():
                                if section_content_list:
                                    with st.expander(f"**{section_name}**"):
                                        for item in section_content_list:
                                            # Assuming section content items might be lists or paragraphs
                                            if isinstance(item, list): # For bullet points
                                                for sub_item in item:
                                                    st.markdown(f"- {sub_item}")
                                            else: # For paragraphs
                                                st.write(item)
                        elif block["type"] == "paragraph" and block["content"]:
                            st.write("\n".join(block["content"]))
                # --- FIN DE LA MODIFICACIÓN ---
                else: # Fallback for unparsed or other types
                    st.write(st.session_state.raw_generated_text_for_display) # Simple display if parsing fails or not defined

                st.info(f"El contenido del {st.session_state.generated_doc_display_type.capitalize()} se ha generado y está listo para descargar como **{st.session_state.generated_doc_display_name}**.")

                # Botón de descarga simple, usando los bytes ya preparados
                st.download_button(
                    label=f"📥 Descargar {st.session_state.generated_doc_display_name}",
                    data=st.session_state.generated_doc_display_content,
                    file_name=st.session_state.generated_doc_display_name,
                    mime=st.session_state.generated_doc_display_mime,
                    key=f"download_generated_doc_display_{st.session_state.generated_doc_display_type}",
                )
        else:
            st.warning("⚠️ Por favor, sube un archivo de temario válido para generar contenido.")

elif st.session_state.selected_section == "🔢 Problemas Matemáticos":
    st.header("🔢 Generador de Problemas Matemáticos")
    st.write("Genera problemas de matemáticas personalizados con soluciones paso a paso. Puedes introducir el problema escribiéndolo o subiendo una imagen.")

    # Estado de sesión para el problema generado
    if 'math_problem_output' not in st.session_state:
        st.session_state.math_problem_output = None
    if 'math_problem_pdf_bytes' not in st.session_state:
        st.session_state.math_problem_pdf_bytes = None
    if 'math_problem_file_name' not in st.session_state:
        st.session_state.math_problem_file_name = ""
    if 'math_problem_input_text' not in st.session_state:
        st.session_state.math_problem_input_text = ""
    if 'math_problem_uploaded_image_data' not in st.session_state:
        st.session_state.math_problem_uploaded_image_data = None # Store raw bytes for display

    col1, col2 = st.columns(2)
    with col1:
        math_topic = st.selectbox(
            "Tema de Matemáticas:",
            ["Aritmética", "Álgebra", "Geometría", "Cálculo", "Estadística", "General"],
            key="math_topic_selector"
        )
    with col2:
        math_level = st.selectbox(
            "Nivel Educativo:",
            ["Primaria", "Secundaria", "Universidad"],
            key="math_level_selector"
        )

    st.markdown("---")
    st.subheader("Entrada del Problema")

    # Text input for manual problem description or LaTeX code
    st.session_state.math_problem_input_text = st.text_area(
        "**1. Describe el problema o pega aquí su código LaTeX:**",
        value=st.session_state.math_problem_input_text, # Retain value across reruns
        placeholder="Ej. Resuelve $x^2 - 5x + 6 = 0$. \nO escribe: 'Crea un problema de la vida real sobre interés compuesto para secundaria.'",
        height=150,
        key="math_text_input"
    )

    # File uploader for image input
    uploaded_math_image = st.file_uploader(
        "**2. O sube una imagen del problema (PNG, JPG, JPEG):**",
        type=["png", "jpg", "jpeg"],
        key="math_image_uploader"
    )

    if uploaded_math_image:
        st.session_state.math_problem_uploaded_image_data = uploaded_math_image.getvalue() # Store raw bytes
        st.image(uploaded_math_image, caption="Imagen del problema cargada", width=200)
        st.warning("Procesando imagen con OCR. La precisión puede variar con la complejidad de la escritura o la calidad de la imagen.")
    elif 'math_problem_uploaded_image_data' in st.session_state and st.session_state.math_problem_uploaded_image_data is not None:
        # If an image was previously uploaded and still in state, display it.
        # This uses BytesIO to convert raw bytes back to a file-like object for st.image
        st.image(BytesIO(st.session_state.math_problem_uploaded_image_data), caption="Imagen del problema previamente cargada", width=200)

    st.markdown("""
        ---
        **💡 Consejos para la Entrada de Problemas:**
        - **Texto:** Puedes escribir el problema directamente o pegar **código LaTeX** para ecuaciones complejas.
        - **LaTeX:** Si el problema es complejo, usa un editor de LaTeX online para escribirlo visualmente y luego copia el código generado.
          - [Overleaf Equation Editor](https://www.overleaf.com/learn/latex/Equations#The_equation_environment)
          - [CodeCogs Online LaTeX Equation Editor](https://www.codecogs.com/latex/eqneditor.php)
    """)
    st.markdown("---")

    if st.button("Generar Problema Matemático", key="generate_math_problem_btn"):
        if not openai.api_key:
            st.error("La clave de API de OpenAI no está configurada. Por favor, revisa la configuración del script.")
        else:
            problem_text_for_ai = st.session_state.math_problem_input_text

            # If an image was uploaded, perform OCR and append its text
            if st.session_state.math_problem_uploaded_image_data:
                with st.spinner("Realizando OCR en la imagen..."):
                    try:
                        # Convert bytes to PIL Image for Tesseract
                        image_from_bytes = Image.open(BytesIO(st.session_state.math_problem_uploaded_image_data))
                        # For better OCR, sometimes converting to grayscale and increasing DPI can help
                        # image_from_bytes = image_from_bytes.convert('L') # Convert to grayscale
                        # image_from_bytes.info['dpi'] = (300, 300) # Set DPI if needed

                        ocr_text = pytesseract.image_to_string(image_from_bytes, lang="spa")
                        if ocr_text.strip():
                            # Append OCR text to manual input, clearly separating them
                            if problem_text_for_ai.strip():
                                problem_text_for_ai += f"\n\n--- Texto extraído de la imagen (OCR):\n{ocr_text}\n---"
                            else: # If only image was provided
                                problem_text_for_ai = ocr_text
                            st.success("Texto de la imagen extraído con OCR.")
                        else:
                            st.warning("No se pudo extraer texto significativo de la imagen con OCR. Por favor, intenta subir una imagen más clara o ingresa el problema manualmente.")
                    except Exception as e:
                        st.error(f"Error al procesar la imagen con OCR: {e}. Asegúrate de que Tesseract esté bien instalado y la imagen sea clara.")
                        st.session_state.math_problem_uploaded_image_data = None # Clear bad image data

            if not problem_text_for_ai.strip():
                st.error("Por favor, describe el problema en el cuadro de texto o sube una imagen.")
                st.session_state.math_problem_output = None
                st.session_state.math_problem_pdf_bytes = None
                st.stop() # Stop execution here if no input

            with st.spinner("⏳ Generando problema y solución con IA..."):
                try:
                    prompt = f"""
                    Eres un experto en matemáticas y un pedagogo. Genera un problema de matemáticas de {math_topic} para nivel {math_level}.
                    El usuario ha proporcionado la siguiente descripción/ejemplo de problema. Si hay un problema explícito, resuélvelo. Si es una instrucción para crear un problema, genera uno nuevo basado en las indicaciones.
                    Asegúrate de que todas las expresiones matemáticas en el problema y en los pasos de la solución estén escritas usando **notación LaTeX in-line (encerrada entre $) o de bloque (encerrada entre $$)** para una visualización clara en el PDF.

                    Sigue este formato exacto:
                    Problema: [Aquí va el problema de matemáticas, con expresiones matemáticas en LaTeX]

                    Pasos:
                    1. [Paso 1 de la solución, con expresiones matemáticas en LaTeX]
                    2. [Paso 2 de la solución, con expresiones matemáticas en LaTeX]
                    ...
                    N. [Último paso de la solución, con expresiones matemáticas en LaTeX]

                    Respuesta Final: [Aquí va la respuesta final, con expresiones matemáticas en LaTeX]

                    Descripción/Instrucción/Problema del usuario:
                    {problem_text_for_ai}
                    """

                    response = openai.chat.completions.create(
                        model="gpt-4o-mini",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.7,
                    )
                    raw_math_problem_text = response.choices[0].message.content
                    st.session_state.math_problem_raw_text = raw_math_problem_text

                    parsed_data = parse_and_format_math_problem(raw_math_problem_text)
                    st.session_state.math_problem_output = parsed_data

                    # Generate PDF
                    pdf_bytes = generate_pdf_from_math_problem(parsed_data)
                    st.session_state.math_problem_pdf_bytes = pdf_bytes

                    file_name_prefix = math_topic.replace(" ", "_").lower()
                    st.session_state.math_problem_file_name = f"Problema_Matematicas_{file_name_prefix}.pdf"

                    # Automatically save to local files history
                    if save_file_to_local_disk_and_update_metadata(
                        pdf_bytes,
                        f"Problema_Matematicas_{file_name_prefix}",
                        ".pdf",
                        "problema_matematicas_generado"
                    ):
                        st.success("✅ Problema generado y guardado localmente como PDF.")
                        st.info("Puedes descargar el archivo o revisarlo en 'Mis Archivos Guardados (Local)'.")
                    else:
                        st.error("No se pudo guardar el archivo localmente. La descarga no estará disponible.")

                except openai.RateLimitError:
                    st.error("🚫 ¡Error de API! Se han agotado los tokens de OpenAI. Por favor, verifica tu plan y facturación en platform.openai.com/usage.")
                except openai.APIError as e:
                    st.error(f"🚫 Error de la API de OpenAI: {e}. Por favor, inténtalo de nuevo más tarde.")
                except Exception as e:
                    st.error(f"Ocurrió un error inesperado al generar el problema: {e}.")
            st.rerun() # Rerun to display the generated problem/PDF download button

    if st.session_state.math_problem_output:
        st.markdown("### 📝 Problema Generado y Solución:")
        st.markdown(f"**Problema:** {st.session_state.math_problem_output['problem']}")

        with st.expander("Ver Solución Paso a Paso"):
            for i, step in enumerate(st.session_state.math_problem_output['steps']):
                st.markdown(f"**{i+1}.** {step}")
            st.markdown(f"**Respuesta Final:** {st.session_state.math_problem_output['answer']}")

        if st.session_state.math_problem_pdf_bytes:
            st.download_button(
                label="📥 Descargar Solución en PDF",
                data=st.session_state.math_problem_pdf_bytes,
                file_name=st.session_state.math_problem_file_name,
                mime="application/pdf",
                key="download_math_problem_pdf"
            )
    else:
        st.info("Ingresa los detalles y haz clic en 'Generar Problema Matemático' para empezar.")


elif st.session_state.selected_section == "💬 Chat Educativo":
    st.header("💬 Chat Educativo con IA")

    # Inicializar estado para la personalidad y el nivel
    if "nivel_docente_chat" not in st.session_state:
        st.session_state.nivel_docente_chat = "Primaria"
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "ai_persona" not in st.session_state:
        st.session_state.ai_persona = "🎓 Asistente Pedagógico"
    if "greeted_current_config" not in st.session_state:
        st.session_state.greeted_current_config = False
    # Nuevo estado para controlar la visibilidad del cargador de archivos
    if 'show_file_upload_options_chat' not in st.session_state:
        st.session_state.show_file_upload_options_chat = False
    # Estado para almacenar el archivo subido temporalmente antes de enviarlo con el prompt
    if 'uploaded_file_info_pending_send' not in st.session_state:
        st.session_state.uploaded_file_info_pending_send = None


    # Guardar estado actual antes de permitir cambios
    previous_level = st.session_state.nivel_docente_chat
    previous_persona = st.session_state.ai_persona

    # 1. Selector de Personalidad (SIEMPRE PRIMERO)
    new_persona_options = ["🎓 Asistente Pedagógico",
                           "🧑‍🎓 Guía Escolar",
                           "💬 Asistente General"]
    new_persona = st.radio(
        "Selecciona la **Personalidad del Asistente IA**:",
        new_persona_options,
        index=new_persona_options.index(st.session_state.ai_persona),
        key="selector_persona_chat"
    )

    # Actualizar la personalidad de la IA en el estado de sesión
    st.session_state.ai_persona = new_persona

    # 2. Selector de Nivel Educativo (CONDICIONAL)
    # Solo mostrar el selector de nivel si la personalidad NO es "Asistente General"
    if st.session_state.ai_persona != "💬 Asistente General":
        new_level = st.selectbox(
            "Selecciona el **Nivel Educativo** (para la complejidad de las respuestas):",
            ["Preescolar", "Primaria", "Secundaria", "Universidad"],
            index=["Preescolar", "Primaria", "Secundaria", "Universidad"].index(st.session_state.nivel_docente_chat),
            key="selector_nivel_chat_dynamic"
        )
        st.session_state.nivel_docente_chat = new_level # Actualizar nivel en estado de sesión
    else:
        new_level = st.session_state.nivel_docente_chat # Keep the current value, but it's not used contextually

    # Verificar si la configuración ha cambiado para resetear el chat y dar un nuevo saludo
    if new_level != previous_level or new_persona != previous_persona:
        st.session_state.chat_history = []
        st.session_state.greeted_current_config = False
        st.session_state.uploaded_file_info_pending_send = None # Clear any pending uploads
        st.rerun()

    # Mostrar el mensaje de info actualizado
    if st.session_state.ai_persona == "💬 Asistente General":
        st.info(f"La IA actúa como **{st.session_state.ai_persona}**.")
    else:
        st.info(f"La IA actúa como **{st.session_state.ai_persona}** para el nivel **{st.session_state.nivel_docente_chat}**.")

    # Saludo inicial o al cambiar configuración
    if not st.session_state.greeted_current_config:
        saludos_pedagogico = {
            "Preescolar": "¡Hola, colega de Preescolar! 👋 Como Asistente Pedagógico, estoy aquí para brindarte apoyo y recursos educativos. ¿En qué puedo asistirte hoy?",
            "Primaria": "¡Saludos, docente de Primaria! 👋 Estoy listo para ofrecerte insights pedagógicos y estrategias de enseñanza para este nivel. ¿Cómo puedo colaborar contigo?",
            "Secundaria": "¡Bienvenido, Asistente Pedagógico! 👋 Mi objetivo es proporcionarte análisis profundos y soluciones para el ámbito educativo secundario. ¿Qué necesitas?",
            "Universidad": "¡Hola, profesional académico! 👋 Como Asistente Pedagógico, estoy a tu disposición para debatir y ofrecer recursos avanzados en temas universitarios. ¿En qué puedo servirte!"
        }
        saludos_guia_escolar = {
            "Preescolar": "¡Hola, amigo! 👋 Soy tu Guía Escolar. ¡Vamos a aprender y divertirnos! ¿Qué quieres que te explique hoy?",
            "Primaria": "¡Hola! 👋 Soy tu Guía Escolar. ¿Necesitas ayuda con la tarea o tienes curiosidad sobre algo? ¡Pregúntame lo que sea!",
            "Secundaria": "¡Qué onda! 👋 Soy tu Guía Escolar. Estoy aquí para hacer que los temas difíciles sean más fáciles de entender. ¿Cuál es tu duda?",
            "Universidad": "¡Saludos! 👋 Soy tu Guía Escolar. Te ofrezco explicaciones claras y apoyo en tus estudios. ¿En qué puedo echarte una mano!"
        }
        saludos_general = "¡Hola! 👋 Soy tu Asistente Conversacional. Estoy aquí para conversar sobre una amplia variedad de temas, incluyendo emociones, hobbies, información general, consejos profesionales, cultura y más. ¿De qué te gustaría charlar hoy?"

        if st.session_state.ai_persona == "🎓 Asistente Pedagógico":
            greeting = saludos_pedagogico[st.session_state.nivel_docente_chat]
        elif st.session_state.ai_persona == "🧑‍🎓 Guía Escolar":
            greeting = saludos_guia_escolar[st.session_state.nivel_docente_chat]
        else: # "💬 Asistente General"
            greeting = saludos_general

        with st.chat_message("assistant"):
            st.write(greeting)
        st.session_state.chat_history.append({"role": "assistant", "content": greeting})
        st.session_state.greeted_current_config = True
        st.rerun()

    chat_history_container = st.container(height=300, border=True)

    with chat_history_container:
        for msg in st.session_state.chat_history:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])
                # Display attached image/document within the chat bubble
                if "image_data" in msg and msg["image_data"]:
                    st.image(
                        f"data:{msg['image_data']['mime_type']};base64,{msg['image_data']['base64']}",
                        caption=f"Adjunto: {msg['image_data']['name']}",
                        width=150
                    )
                if "document_data" in msg and msg["document_data"]:
                    st.info(f"Adjunto: {msg['document_data']['name']}")
                    # You can uncomment below to show a snippet of extracted text
                    # with st.expander("Ver texto extraído del documento"):
                    #    st.code(msg['document_data']['extracted_text'][:500] + "...")

    # Área de entrada de chat y botón de adjuntar
    col_chat_input, col_add_button = st.columns([0.9, 0.1])

    with col_add_button:
        # Botón para mostrar/ocultar el uploader de archivos
        if st.button("➕", key="toggle_upload_button_chat"):
            st.session_state.show_file_upload_options_chat = not st.session_state.show_file_upload_options_chat
            # No rerun here, we want the uploader to just appear/disappear without clearing the input

    # Display the file uploader if the plus button was clicked
    if st.session_state.show_file_upload_options_chat:
        uploaded_file_chat = st.file_uploader(
            "Selecciona un archivo (imagen, PDF, DOCX)",
            type=["png", "jpg", "jpeg", "pdf", "docx"],
            key="chat_file_uploader_universal"
        )
        if uploaded_file_chat:
            # Process the uploaded file and store it in a pending state
            if uploaded_file_chat.type.startswith("image/"):
                base64_img, mime_type_img = get_file_base64_and_mime_type(uploaded_file_chat)
                st.session_state.uploaded_file_info_pending_send = {"type": "image", "base64": base64_img, "mime_type": mime_type_img, "name": uploaded_file_chat.name}
            else: # PDF or DOCX
                extracted_text = extract_text_from_document(uploaded_file_chat)
                st.session_state.uploaded_file_info_pending_send = {
                    "type": "document",
                    "name": uploaded_file_chat.name,
                    "extracted_text": extracted_text,
                    "mime_type": uploaded_file_chat.type
                }

            st.session_state.show_file_upload_options_chat = False # Hide the uploader after selection
            # Do NOT call st.rerun() here. This allows the user to type text before sending.
            st.info(f"Archivo '{uploaded_file_chat.name}' adjuntado. Ahora puedes escribir un mensaje y presionar Enter para enviar.")


    # Check if a file is pending and display a placeholder confirmation message near the input
    if st.session_state.uploaded_file_info_pending_send:
        pending_file = st.session_state.uploaded_file_info_pending_send
        st.markdown(f"**Archivo listo para enviar:** `{pending_file['name']}`")
        if pending_file["type"] == "image":
            st.image(
                f"data:{pending_file['mime_type']};base64,{pending_file['base64']}",
                caption="Imagen adjunta (pendiente)",
                width=80
            )

    with col_chat_input:
        # The prompt_usuario will only be non-empty when the user types and presses enter.
        prompt_usuario = st.chat_input("Escribe tu mensaje...", key="chat_input_text")

    # Procesar el mensaje del usuario (texto + archivo si se adjuntó)
    # This block will now only execute when 'prompt_usuario' changes (user presses Enter)
    if prompt_usuario:
        user_content_parts = []

        # Add text prompt
        user_content_parts.append({"type": "text", "text": prompt_usuario})

        # Add pending file data if exists
        if st.session_state.uploaded_file_info_pending_send:
            uploaded_info = st.session_state.uploaded_file_info_pending_send
            if uploaded_info["type"] == "image":
                user_content_parts.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{uploaded_info['mime_type']};base64,{uploaded_info['base64']}"
                    }
                })
            elif uploaded_info["type"] == "document":
                # Add extracted text to the prompt for the LLM
                user_content_parts.append({"type": "text", "text": f"\n\n--- Contenido del documento '{uploaded_info['name']}':\n{uploaded_info['extracted_text']}\n---"})

        # Append user's message (text and/or file) to chat history for display
        user_message_display_entry = {"role": "user", "content": prompt_usuario}
        if st.session_state.uploaded_file_info_pending_send:
            if st.session_state.uploaded_file_info_pending_send["type"] == "image":
                user_message_display_entry["image_data"] = {"base64": st.session_state.uploaded_file_info_pending_send['base64'], "mime_type": st.session_state.uploaded_file_info_pending_send['mime_type'], "name": st.session_state.uploaded_file_info_pending_send['name']}
            elif st.session_state.uploaded_file_info_pending_send["type"] == "document":
                user_message_display_entry["document_data"] = {"name": st.session_state.uploaded_file_info_pending_send['name'], "extracted_text": st.session_state.uploaded_file_info_pending_send['extracted_text'], "mime_type": st.session_state.uploaded_file_info_pending_send['mime_type']}

        st.session_state.chat_history.append(user_message_display_entry)

        # Clear the pending uploaded file info after it's been processed and added to history
        st.session_state.uploaded_file_info_pending_send = None

        if user_content_parts: # Only proceed if there's actual content to send to API
            with st.spinner("🤖 Generando respuesta..."):
                contenido_chat = ""
                if not openai.api_key:
                    st.error("La clave de API de OpenAI no está configurada. Por favor, revisa la configuración del script.")
                    contenido_chat = "Lo siento, no puedo responder. La clave de API de OpenAI no está configurada."
                else:
                    try:
                        # Build system role prompt based on selected persona and level
                        if st.session_state.ai_persona == "🎓 Asistente Pedagógico":
                            system_role_prompt = f"Eres un docente experto en nivel {st.session_state.nivel_docente_chat}. Responde de forma útil, detallada y pedagógica para un docente o jurado. Tu objetivo es proporcionar información precisa y herramientas para la enseñanza."
                        elif st.session_state.ai_persona == "🧑‍🎓 Guía Escolar":
                            system_role_prompt = f"Eres un guía escolar amigable y claro para estudiantes. Responde de forma sencilla, directa y didáctica, adaptando la explicación al nivel de un estudiante de {st.session_state.nivel_docente_chat}. Fomenta el aprendizaje y la curiosidad. Evita jergas complejas y respuestas robotizadas."
                        else: # "💬 Asistente General"
                            system_role_prompt = "Eres un asistente de IA conversacional y amigable, capaz de hablar sobre una amplia variedad de temas, incluyendo emociones, hobbies, información general, consejos profesionales, cultura y más. Responde de manera natural, informada y empática, manteniendo un tono de conversación abierta y adaptándote al contexto de la pregunta. Evita respuestas que suenen demasiado 'IA' o robóticas."

                        full_messages_for_api = [
                            {"role": "system", "content": system_role_prompt},
                            {"role": "user", "content": user_content_parts}
                        ]

                        respuesta_chat = openai.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=full_messages_for_api,
                            temperature=0.7
                        )
                        contenido_chat = respuesta_chat.choices[0].message.content
                    except openai.RateLimitError:
                        st.error("🚫 ¡Error de API! Se han agotaron los tokens de OpenAI. Por favor, verifica tu plan y facturación en platform.openai.com/usage.")
                        contenido_chat = "Lo siento, hubo un problema con la API (límites de uso). Intenta más tarde."
                    except openai.APIError as e:
                        st.error(f"🚫 Error de la API de OpenAI: {e}. Por favor, inténtalo de nuevo más tarde.")
                        contenido_chat = f"Lo siento, hubo un problema con la API: {e}"
                    except Exception as e:
                        st.error(f"Ocurrió un error inesperado al procesar tu mensaje: {e}. Por favor, inténtalo de nuevo.")
                        contenido_chat = f"Lo siento, ocurrió un error inesperado: {e}"

                if contenido_chat: # Append assistant's response only once to history
                    st.session_state.chat_history.append({"role": "assistant", "content": contenido_chat})
            st.rerun()


elif st.session_state.selected_section == "🔄 Convertidor de Docs":
    st.header("🔄 Convertidor de Documentos")

    st.write("Sube tu archivo y selecciona el formato al que deseas convertirlo.")

    uploaded_file_converter = st.file_uploader("Sube un archivo (PDF o DOCX)", type=["pdf", "docx"], key="converter_uploader")

    # Almacenar el contenido convertido en el estado de sesión
    if 'converted_doc_display_bytes' not in st.session_state:
        st.session_state.converted_doc_display_bytes = None
    if 'converted_doc_display_name' not in st.session_state:
        st.session_state.converted_doc_display_name = ""
    if 'converted_doc_display_type' not in st.session_state:
        st.session_state.converted_doc_display_type = ""
    if 'converted_doc_display_mime' not in st.session_state:
        st.session_state.converted_doc_display_mime = ""

    if uploaded_file_converter is not None:
        file_name = uploaded_file_converter.name

        st.subheader("Opciones de Conversión")

        if file_name.endswith(".pdf"):
            st.markdown("### Convertir PDF a DOCX")
            st.write("Extraerá el texto de tu PDF y lo guardará en un documento de Word.")
            if st.button("Convertir a DOCX", key="convert_pdf_to_docx"):
                with st.spinner("Convirtiendo PDF a DOCX..."):
                    texto_extraido = ""
                    # Seek to the beginning of the file before processing
                    uploaded_file_converter.seek(0)
                    try:
                        with pdfplumber.open(uploaded_file_converter) as pdf:
                            for pagina in pdf.pages:
                                texto_extraido += pagina.extract_text() or ""

                        if not texto_extraido.strip():
                            st.warning("No se encontró texto digital, intentando OCR (puede tardar)...")
                            uploaded_file_converter.seek(0) # Reset pointer
                            pdf_bytes = uploaded_file_converter.getvalue()
                            images = convert_from_bytes(pdf_bytes)
                            for i, image in enumerate(images):
                                texto_extraido += pytesseract.image_to_string(image, lang="spa")
                            st.success("OCR completado.")

                        if texto_extraido.strip():
                            document = docx.Document()
                            document.add_paragraph(texto_extraido)

                            doc_buffer = BytesIO()
                            document.save(doc_buffer)
                            doc_buffer.seek(0)
                            docx_bytes_content = doc_buffer.getvalue()

                            # Guardar inmediatamente en el disco local y actualizar metadatos
                            base_name_for_file = file_name.replace(".pdf", "")
                            if save_file_to_local_disk_and_update_metadata(docx_bytes_content, base_name_for_file, ".docx", "conversion_docx_from_pdf"):
                                st.session_state.converted_doc_display_bytes = docx_bytes_content
                                st.session_state.converted_doc_display_name = f"{base_name_for_file}.docx"
                                st.session_state.converted_doc_display_type = "conversion_docx_from_pdf"
                                st.session_state.converted_doc_display_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                                st.info("Puedes descargar el archivo o revisarlo en 'Mis Archivos Guardados (Local)'.")
                            else:
                                st.error("No se pudo guardar el archivo localmente. La descarga no estará disponible.")
                                st.session_state.converted_doc_display_bytes = None
                        else:
                            st.error("❌ No se pudo extraer texto del PDF para convertirlo a DOCX.")
                            st.session_state.converted_doc_display_bytes = None
                    except Exception as e:
                        st.error(f"Ocurrió un error al convertir PDF a DOCX: {e}")
                        st.session_state.converted_doc_display_bytes = None
                st.rerun()

        elif file_name.endswith(".docx"):
            st.markdown("### Convertir DOCX a PDF")
            st.write("Selecciona cómo quieres convertir tu documento Word a PDF.")
            conversion_type = st.radio(
                "Elige tipo de conversión de DOCX a PDF:",
                ["Mantener formato (requiere MS Word)", "Solo texto (no requiere MS Word)"],
                key="docx_to_pdf_type"
            )

            if st.button("Convertir a PDF", key="convert_docx_to_pdf"):
                with st.spinner(f"Convirtiendo DOCX a PDF ({conversion_type})..."):
                    if conversion_type == "Mantener formato (requiere MS Word)":
                        temp_docx_path = ""
                        temp_pdf_path = ""
                        try:
                            # Ensure the file pointer is at the beginning before reading
                            uploaded_file_converter.seek(0)
                            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
                                temp_docx.write(uploaded_file_converter.read())
                                temp_docx_path = temp_docx.name

                            # Define temp_pdf_path before calling convert
                            temp_pdf_fd, temp_pdf_path = tempfile.mkstemp(suffix=".pdf")
                            os.close(temp_pdf_fd) # Close the file descriptor immediately

                            docx_to_pdf_convert(temp_docx_path, temp_pdf_path)

                            with open(temp_pdf_path, "rb") as pdf_file:
                                pdf_bytes_content = pdf_file.read()

                            # Guardar inmediatamente en el disco local y actualizar metadatos
                            base_name_for_file = file_name.replace(".docx", "")
                            if save_file_to_local_disk_and_update_metadata(pdf_bytes_content, base_name_for_file, ".pdf", "conversion_pdf_from_docx_formatted"):
                                st.session_state.converted_doc_display_bytes = pdf_bytes_content
                                st.session_state.converted_doc_display_name = f"{base_name_for_file}.pdf"
                                st.session_state.converted_doc_display_type = "conversion_pdf_from_docx_formatted"
                                st.session_state.converted_doc_display_mime = "application/pdf"
                                st.success("✅ DOCX convertido a PDF (con formato) exitosamente.")
                                st.info("Puedes descargar el archivo o revisarlo en 'Mis Archivos Guardados (Local)'.")
                            else:
                                st.error("No se pudo guardar el archivo localmente. La descarga no estará disponible.")
                                st.session_state.converted_doc_display_bytes = None

                        except FileNotFoundError:
                            st.error("🚫 ¡Error! Parece que Microsoft Word no está instalado o no se encontró. Necesitas Word para esta opción de conversión.")
                            st.session_state.converted_doc_display_bytes = None
                        except Exception as e:
                            st.error(f"Ocurrió un error al convertir DOCX a PDF (con formato): {e}")
                            st.session_state.converted_doc_display_bytes = None
                        finally:
                            if os.path.exists(temp_docx_path):
                                os.remove(temp_docx_path)
                            if 'temp_pdf_path' in locals() and os.path.exists(temp_pdf_path):
                                os.remove(temp_pdf_path)
                        st.rerun()

                    else: # Solo texto (no requiere MS Word)
                        texto_doc = ""
                        # Ensure the file pointer is at the beginning before reading
                        uploaded_file_converter.seek(0)
                        try:
                            doc = docx.Document(uploaded_file_converter)
                            full_text = []
                            for para in doc.paragraphs:
                                full_text.append(para.text)
                            texto_doc = "\n".join(full_text)

                            pdf_buffer = BytesIO()
                            doc_pdf = SimpleDocTemplate(pdf_buffer, pagesize=letter)
                            styles = getSampleStyleSheet()
                            # Use keepWithNext=True for paragraphs to prevent orphaned headings
                            story = [Paragraph(texto_doc.replace("\n", "<br/>"), styles["Normal"])]

                            doc_pdf.build(story)
                            pdf_buffer.seek(0)
                            pdf_bytes_content = pdf_buffer.getvalue()

                            # Guardar inmediatamente en el disco local y actualizar metadatos
                            base_name_for_file = file_name.replace(".docx", "_text")
                            if save_file_to_local_disk_and_update_metadata(pdf_bytes_content, base_name_for_file, ".pdf", "conversion_pdf_from_docx_text"):
                                st.session_state.converted_doc_display_bytes = pdf_bytes_content
                                st.session_state.converted_doc_display_name = f"{base_name_for_file}.pdf"
                                st.session_state.converted_doc_display_type = "conversion_pdf_from_docx_text"
                                st.session_state.converted_doc_display_mime = "application/pdf"
                                st.success("✅ DOCX convertido a PDF (solo texto) exitosamente.")
                                st.info("Puedes descargar el archivo o revisarlo en 'Mis Archivos Guardados (Local)'.")
                            else:
                                st.error("No se pudo guardar el archivo localmente. La descarga no estará disponible.")
                                st.session_state.converted_doc_display_bytes = None

                        except Exception as e:
                            st.error(f"Ocurrió un error al convertir DOCX a PDF (solo texto): {e}")
                            st.session_state.converted_doc_display_bytes = None
                        st.rerun()
        else:
            st.warning("⚠️ Formato de archivo no soportado para conversión. Por favor, sube un archivo PDF o DOCX.")

        # Mostrar el botón de descarga si hay contenido convertido
        if st.session_state.converted_doc_display_bytes is not None:
            st.markdown("### ✅ Resultado de Conversión Disponible")

            st.download_button(
                label=f"📥 Descargar {st.session_state.converted_doc_display_name}",
                data=st.session_state.converted_doc_display_bytes,
                file_name=st.session_state.converted_doc_display_name,
                mime=st.session_state.converted_doc_display_mime,
                key=f"download_converted_doc_display_{st.session_state.converted_doc_display_type}",
            )

# ==================== NUEVA SECCIÓN: JUEGO EDUCATIVO ====================
elif st.session_state.selected_section == "🎮 Juegos Educativos":
    st.header("🎮 Juego Educativo: ¡Aprende y Diviértete!")
    st.write("¡Elige un juego para poner a prueba tus conocimientos!")

    # Solo queda la opción de Math Slither
    game_selection = st.radio(
        "Selecciona el tipo de juego:",
        ["Juego Externo: Math Slither"],
        key="game_type_selection"
    )

    if game_selection == "Juego Externo: Math Slither":
        st.subheader("Juego Externo: Math Slither")
        st.write("A continuación se intentará cargar el juego 'Math Slither'.")

        # URL del juego Math Slither en mathgames.com
        math_slither_url = "https://www.mathgames.com/play/math-slither.html"

        # Puedes ajustar el ancho (width) y el alto (height) para que se vea bien.
        components.html(
            f"""
            <iframe src="{math_slither_url}" width="700" height="500" frameborder="0" allowfullscreen></iframe>
            """,
            height=550, # El alto total del componente en Streamlit
        )
        st.info("Si el juego no se muestra, es probable que la página no permita ser incrustada. En ese caso, la mejor opción sería que los alumnos accedan al juego directamente desde el navegador.")


# ==================== SECCIÓN: MIS ARCHIVOS GUARDADOS (LOCAL) ====================
elif st.session_state.selected_section == "🗄️ Mis Archivos":
    st.header("🗄️ Mis Archivos Guardados (Local)")
    st.write("Aquí verás los documentos que has guardado en la carpeta local de tu aplicación.")

    if st.session_state.local_files_history:
        st.subheader("Documentos Guardados:")
        # Usar la lista minimalista directamente
        for i, file_data in enumerate(st.session_state.local_files_history):
            file_name = file_data.get('name', f"descarga_sin_nombre_{i}.txt")
            file_type = file_data.get('type', 'desconocido')
            generated_at = file_data.get('generated_at', 'Fecha desconocida')

            # Usar columnas para alinear elementos
            col_name, col_date, col_download, col_delete = st.columns([0.5, 0.25, 0.15, 0.1])

            with col_name:
                st.write(f"**{file_name}**")
            with col_date:
                st.markdown(f"<p style='font-size: 0.9em; color: gray;'>{generated_at}</p>", unsafe_allow_html=True)
            with col_download:
                download_label = "Descargar"
                download_file_path = file_data.get('local_path')

                mime_type = "application/octet-stream"
                if file_name.lower().endswith(".txt"):
                    mime_type = "text/plain"
                elif file_name.lower().endswith(".docx"):
                    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                elif file_name.lower().endswith(".pdf"):
                    mime_type = "application/pdf"

                if os.path.exists(download_file_path):
                    try:
                        with open(download_file_path, "rb") as f:
                            file_bytes = f.read()
                        st.download_button(
                            label=download_label,
                            data=file_bytes,
                            file_name=file_name,
                            mime=mime_type,
                            key=f"download_saved_local_file_{file_data.get('id')}"
                        )
                    except Exception as e:
                        st.error(f"No se pudo leer el archivo '{file_name}' para descargar: {e}")
                else:
                    st.warning(f"Archivo no encontrado en el disco: {file_name}")

            with col_delete:
                # Need to give a unique key to each delete button
                if st.button("Eliminar", key=f"delete_saved_local_file_{file_data.get('id')}"):
                    delete_local_file(file_data.get('id'))
            st.markdown("---")
    else:
        st.info("No tienes archivos guardados en tu carpeta local. ¡Genera o convierte algunos en las otras secciones y guárdalos!")
